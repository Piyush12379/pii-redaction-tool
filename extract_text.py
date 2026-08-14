from docx import Document


def clean_text(text):
    """
    Clean unnecessary whitespace from extracted text.
    """
    return " ".join(text.split())


def extract_document_text(file_path):
    """
    Extract text from paragraphs and tables.
    Returns a list of unique text blocks.
    """

    document = Document(file_path)

    text_blocks = []

    # Extract paragraphs
    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text)

        if text:
            text_blocks.append(text)

    # Extract tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = clean_text(cell.text)

                if text:
                    text_blocks.append(text)

    # Remove exact duplicates
    unique_blocks = list(dict.fromkeys(text_blocks))

    return unique_blocks


if __name__ == "__main__":

    file_path = "Prospectus.docx"

    blocks = extract_document_text(file_path)

    print("Total unique text blocks:", len(blocks))

    print("\nFirst 20 text blocks:\n")

    for block in blocks[:20]:
        print(block)
