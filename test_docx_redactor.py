import os

from docx import Document

from docx_redactor import redact_docx


# ============================================================
# FILE PATHS
# ============================================================

INPUT_FILE = "Prospectus.docx"

OUTPUT_FILE = "Redacted_Prospectus.docx"


# ============================================================
# RUN REDACTION
# ============================================================

print("=" * 80)
print("PHASE 11 - DOCX REDACTION")
print("=" * 80)

print()
print("Input :", INPUT_FILE)
print("Output:", OUTPUT_FILE)

print()
print("Starting redaction...")


total_detected = redact_docx(
    INPUT_FILE,
    OUTPUT_FILE
)


# ============================================================
# VERIFY OUTPUT
# ============================================================

print()
print("=" * 80)
print("REDACTION COMPLETE")
print("=" * 80)

print(
    "PII occurrences processed:",
    total_detected
)

print(
    "Output file:",
    OUTPUT_FILE
)


# ============================================================
# CHECK FILE EXISTS
# ============================================================

if os.path.exists(OUTPUT_FILE):

    file_size = os.path.getsize(
        OUTPUT_FILE
    )

    print(
        "Output size:",
        file_size,
        "bytes"
    )

else:

    print(
        "ERROR: Output file was not created."
    )

    raise SystemExit(1)


# ============================================================
# CHECK DOCX CAN BE OPENED
# ============================================================

print()
print("Testing generated DOCX...")


try:

    document = Document(
        OUTPUT_FILE
    )

    paragraph_count = len(
        document.paragraphs
    )

    table_count = len(
        document.tables
    )

    print(
        "DOCX opened successfully."
    )

    print(
        "Paragraphs:",
        paragraph_count
    )

    print(
        "Tables:",
        table_count
    )

except Exception as error:

    print(
        "ERROR: Generated DOCX could not be opened."
    )

    print(error)

    raise SystemExit(1)


print()
print("=" * 80)
print("PHASE 11 TEST PASSED")
print("=" * 80)
