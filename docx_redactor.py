import os
from copy import deepcopy

from docx import Document

from pii_detector import detect_all_pii
from redaction_engine import redact_text


# ============================================================
# REDACT A PARAGRAPH
# ============================================================

def redact_paragraph(paragraph):
    """
    Detect and redact PII inside a DOCX paragraph.

    The paragraph text is replaced while preserving the
    paragraph itself and its surrounding document structure.
    """

    original_text = paragraph.text

    if not original_text.strip():
        return 0

    pii_results = detect_all_pii(
        original_text
    )

    total_detected = sum(
        len(values)
        for values in pii_results.values()
    )

    if total_detected == 0:
        return 0

    redacted_text = redact_text(
        original_text,
        pii_results
    )

    if redacted_text == original_text:
        return 0

    # --------------------------------------------------------
    # Remove existing runs.
    #
    # We keep the paragraph itself, but replace its textual
    # contents with the redacted version.
    # --------------------------------------------------------

    for run in paragraph.runs:
        run.text = ""

    if paragraph.runs:
        paragraph.runs[0].text = redacted_text
    else:
        paragraph.add_run(
            redacted_text
        )

    return total_detected


# ============================================================
# REDACT TABLE CELL
# ============================================================

def redact_cell(cell):
    """
    Redact all paragraphs inside a table cell.
    """

    total = 0

    for paragraph in cell.paragraphs:

        total += redact_paragraph(
            paragraph
        )

    # --------------------------------------------------------
    # Nested tables
    # --------------------------------------------------------

    for table in cell.tables:

        total += redact_table(
            table
        )

    return total


# ============================================================
# REDACT TABLE
# ============================================================

def redact_table(table):
    """
    Redact all text contained in a table.
    """

    total = 0

    for row in table.rows:

        for cell in row.cells:

            total += redact_cell(
                cell
            )

    return total


# ============================================================
# REDACT HEADER
# ============================================================

def redact_header(header):
    """
    Redact paragraphs and tables in a document header.
    """

    total = 0

    for paragraph in header.paragraphs:

        total += redact_paragraph(
            paragraph
        )

    for table in header.tables:

        total += redact_table(
            table
        )

    return total


# ============================================================
# REDACT FOOTER
# ============================================================

def redact_footer(footer):
    """
    Redact paragraphs and tables in a document footer.
    """

    total = 0

    for paragraph in footer.paragraphs:

        total += redact_paragraph(
            paragraph
        )

    for table in footer.tables:

        total += redact_table(
            table
        )

    return total


# ============================================================
# REDACT COMPLETE DOCX
# ============================================================

def redact_docx(
    input_path,
    output_path
):
    """
    Read a DOCX, detect PII, redact it, and save a new DOCX.

    The original input document is never modified.

    Returns:
        total number of detected PII occurrences.
    """

    if not os.path.exists(input_path):

        raise FileNotFoundError(
            f"Input file not found: {input_path}"
        )

    document = Document(
        input_path
    )

    total_detected = 0

    # ========================================================
    # MAIN DOCUMENT PARAGRAPHS
    # ========================================================

    for paragraph in document.paragraphs:

        total_detected += redact_paragraph(
            paragraph
        )

    # ========================================================
    # MAIN DOCUMENT TABLES
    # ========================================================

    for table in document.tables:

        total_detected += redact_table(
            table
        )

    # ========================================================
    # HEADERS AND FOOTERS
    # ========================================================

    for section in document.sections:

        total_detected += redact_header(
            section.header
        )

        total_detected += redact_footer(
            section.footer
        )

        # First-page header/footer
        if section.different_first_page_header_footer:

            total_detected += redact_header(
                section.first_page_header
            )

            total_detected += redact_footer(
                section.first_page_footer
            )

        # Even-page header/footer
        if section.even_page_header.is_linked_to_previous is False:

            total_detected += redact_header(
                section.even_page_header
            )

            total_detected += redact_footer(
                section.even_page_footer
            )

    # ========================================================
    # SAVE OUTPUT
    # ========================================================

    document.save(
        output_path
    )

    return total_detected
