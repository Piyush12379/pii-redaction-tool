from docx import Document

document = Document("Prospectus.docx")

print("Document loaded successfully!")
print("Number of paragraphs:", len(document.paragraphs))
print("Number of tables:", len(document.tables))

print("\n--- FIRST 10 NON-EMPTY PARAGRAPHS ---")

count = 0

for paragraph in document.paragraphs:
    text = paragraph.text.strip()

    if text:
        print(text)
        count += 1

    if count == 10:
        break


print("\n--- FIRST TABLE ---")

if document.tables:

    table = document.tables[0]

    for row in table.rows:

        row_data = []

        for cell in row.cells:
            row_data.append(cell.text.strip())

        print(row_data)